import io
import uuid
from datetime import UTC, datetime

from httpx import ASGITransport, AsyncClient
from minio import Minio
from sqlalchemy import select

from core.config import settings
from core.db import SessionLocal
from core.main import app
from core.models import Document, User


async def _seed_owner() -> uuid.UUID:
    owner_id = uuid.uuid4()
    async with SessionLocal() as db:
        db.add(
            User(
                id=owner_id,
                email=f"owner-{owner_id.hex}@example.com",
                display_name="Owner",
                status="active",
                created_at=datetime.now(UTC),
            )
        )
        await db.commit()
    return owner_id


def _minio() -> Minio:
    return Minio(
        settings.minio_endpoint,
        access_key=settings.minio_access_key,
        secret_key=settings.minio_secret_key,
        secure=settings.minio_use_ssl,
    )


async def _seed_document(
    document_id: uuid.UUID,
    storage_key: str,
    mime: str,
    content: bytes,
    file_name: str,
) -> None:
    client = _minio()
    client.put_object(
        settings.minio_bucket,
        storage_key,
        io.BytesIO(content),
        length=len(content),
        content_type=mime,
    )

    owner_id = await _seed_owner()
    async with SessionLocal() as db:
        db.add(
            Document(
                id=document_id,
                title=file_name,
                body="",
                created_at=datetime.now(UTC),
                file_name=file_name,
                mime_type=mime,
                size_bytes=len(content),
                storage_key=storage_key,
                file_hash=uuid.uuid4().hex,
                ingestion_status="Pending",
                ingestion_error=None,
                owner_id=owner_id,
            )
        )
        await db.commit()


async def _load_document(document_id: uuid.UUID) -> Document:
    async with SessionLocal() as db:
        result = await db.execute(select(Document).where(Document.id == document_id))
        return result.scalar_one()


SERVICE_HEADERS = {"X-Service-Token": settings.service_token}


async def test_post_ingest_requires_service_token() -> None:
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
        response = await ac.post(f"/ingest/{uuid.uuid4()}")
    assert response.status_code == 401


async def test_post_ingest_extracts_txt_and_marks_ready() -> None:
    document_id = uuid.uuid4()
    storage_key = f"test-ingest/{uuid.uuid4().hex}.txt"
    await _seed_document(document_id, storage_key, "text/plain", b"Hello world", "hello.txt")

    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
        response = await ac.post(f"/ingest/{document_id}", headers=SERVICE_HEADERS)

    assert response.status_code == 204

    doc = await _load_document(document_id)
    assert doc.body == "Hello world"
    assert doc.ingestion_status == "Ready"


async def test_post_ingest_returns_404_when_document_missing() -> None:
    missing_id = uuid.uuid4()
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
        response = await ac.post(f"/ingest/{missing_id}", headers=SERVICE_HEADERS)
    assert response.status_code == 404


async def test_post_ingest_extracts_markdown() -> None:
    document_id = uuid.uuid4()
    storage_key = f"test-ingest/{uuid.uuid4().hex}.md"
    content = b"# Heading\n\nSome **bold** text."
    await _seed_document(document_id, storage_key, "text/markdown", content, "doc.md")

    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
        response = await ac.post(f"/ingest/{document_id}", headers=SERVICE_HEADERS)
    assert response.status_code == 204

    doc = await _load_document(document_id)
    assert doc.body == content.decode("utf-8")
    assert doc.ingestion_status == "Ready"


async def test_post_ingest_marks_failed_when_blob_missing() -> None:
    document_id = uuid.uuid4()
    owner_id = await _seed_owner()
    # Seed DB row but NOT MinIO
    async with SessionLocal() as db:
        db.add(
            Document(
                id=document_id,
                title="ghost.txt",
                body="",
                created_at=datetime.now(UTC),
                file_name="ghost.txt",
                mime_type="text/plain",
                size_bytes=10,
                storage_key=f"test-ingest/missing-{uuid.uuid4().hex}.txt",
                file_hash=uuid.uuid4().hex,
                ingestion_status="Pending",
                ingestion_error=None,
                owner_id=owner_id,
            )
        )
        await db.commit()

    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
        response = await ac.post(f"/ingest/{document_id}", headers=SERVICE_HEADERS)
    assert response.status_code == 204

    doc = await _load_document(document_id)
    assert doc.ingestion_status == "Failed"
    assert doc.ingestion_error is not None and doc.ingestion_error != ""


async def test_post_ingest_extracts_docx() -> None:
    from docx import Document as DocxDocument

    docx_doc = DocxDocument()
    docx_doc.add_paragraph("Hello DOCX paragraph")
    buf = io.BytesIO()
    docx_doc.save(buf)
    content = buf.getvalue()

    document_id = uuid.uuid4()
    storage_key = f"test-ingest/{uuid.uuid4().hex}.docx"
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    await _seed_document(document_id, storage_key, mime, content, "doc.docx")

    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
        response = await ac.post(f"/ingest/{document_id}", headers=SERVICE_HEADERS)
    assert response.status_code == 204

    doc = await _load_document(document_id)
    assert "Hello DOCX paragraph" in doc.body
    assert doc.ingestion_status == "Ready"


async def test_post_ingest_extracts_pdf() -> None:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    pdf.cell(0, 10, "Hello PDF body")
    content = bytes(pdf.output())

    document_id = uuid.uuid4()
    storage_key = f"test-ingest/{uuid.uuid4().hex}.pdf"
    await _seed_document(document_id, storage_key, "application/pdf", content, "doc.pdf")

    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
        response = await ac.post(f"/ingest/{document_id}", headers=SERVICE_HEADERS)
    assert response.status_code == 204

    doc = await _load_document(document_id)
    assert "Hello PDF body" in doc.body
    assert doc.ingestion_status == "Ready"
